"""Office file extraction utilities for DOCX and XLSX files."""

import os
from typing import Any


def extract_docx(path: str) -> dict[str, Any]:
    """Extract content from a .docx file.

    Args:
        path: Path to the .docx file.

    Returns:
        dict with keys: text, tables, source_file, file_type.
    """
    from docx import Document

    doc = Document(path)

    # Extract paragraphs
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)

    # Extract tables
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            rows.append(cells)
        tables.append(rows)

    return {
        "text": text,
        "tables": tables,
        "source_file": os.path.basename(path),
        "file_type": "docx",
    }


def extract_xlsx(path: str) -> dict[str, Any]:
    """Extract content from an .xlsx file.

    Args:
        path: Path to the .xlsx file.

    Returns:
        dict with keys: sheets, source_file, file_type.
    """
    from openpyxl import load_workbook

    wb = load_workbook(path, read_only=True, data_only=True)

    sheets = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            # Convert all cell values to strings or empty strings
            row_data = [str(cell) if cell is not None else "" for cell in row]
            rows.append(row_data)
        sheets.append({"name": sheet_name, "rows": rows})

    wb.close()

    return {
        "sheets": sheets,
        "source_file": os.path.basename(path),
        "file_type": "xlsx",
    }
